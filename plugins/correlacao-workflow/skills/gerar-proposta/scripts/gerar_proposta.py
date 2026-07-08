#!/usr/bin/env python3
"""
Monta o .docx final da proposta de forma determinística.

Este script NÃO decide o conteúdo da proposta — o texto das seções
criativas já deve vir pronto (escrito e revisado pelo agente nos Passos 3-5
do SKILL.md) e os campos determinísticos (preço, prazo, cláusulas) já
devem vir da tabela de referência da organização, ou como o marcador
"[PENDENTE — preencher com tabela oficial]" enquanto essa tabela não
existir. A única responsabilidade deste script é montar o documento com
formatação consistente — a formatação não deve ser reinventada pelo modelo
a cada proposta.

Uso:
    python3 gerar_proposta.py dados_proposta.json saida.docx [--template modelo.docx]

Onde dados_proposta.json tem o formato:
    {
      "cliente": "Nome do Cliente",
      "data": "02/07/2026",
      "secoes_criativas": {
        "abertura": "texto...",
        "diagnostico": "texto...",
        "proposta_de_valor": "texto..."
      },
      "campos_deterministicos": {
        "escopo": "texto vindo da tabela, ou [PENDENTE — preencher com tabela oficial]",
        "preco": "[PENDENTE — preencher com tabela oficial]",
        "prazo": "[PENDENTE — preencher com tabela oficial]",
        "condicoes_pagamento": "[PENDENTE — preencher com tabela oficial]"
      },
      "templates_usados": ["proposta-123: selecionada por correlação de escopo de automação"]
    }

Se algum campo determinístico vier com o marcador "[PENDENTE", o script
mantém o marcador visível no documento em vez de tentar preenchê-lo — a
ideia é que fique óbvio na revisão humana que aquele valor ainda não foi
confirmado, nunca disfarçado como um valor real.
"""

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PENDENTE_MARCADOR = "[PENDENTE"
ALERTA_COR = RGBColor(0xB0, 0x00, 0x00)

SECAO_TITULOS = {
    "abertura": "Abertura",
    "diagnostico": "Diagnóstico",
    "proposta_de_valor": "Proposta de Valor",
}

CAMPO_TITULOS = {
    "escopo": "Escopo",
    "preco": "Investimento",
    "prazo": "Prazo",
    "condicoes_pagamento": "Condições de Pagamento",
}


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text, alerta=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if alerta:
        run.font.color.rgb = ALERTA_COR
        run.bold = True
    return p


def build_document(dados: dict, template_path: str | None) -> Document:
    doc = Document(template_path) if template_path else Document()

    doc.add_heading(f"Proposta Comercial — {dados.get('cliente', '[cliente não informado]')}", level=0)
    meta = doc.add_paragraph()
    meta.add_run(f"Data: {dados.get('data', '[data não informada]')}").italic = True
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    secoes = dados.get("secoes_criativas", {})
    for chave, titulo in SECAO_TITULOS.items():
        if chave in secoes:
            add_heading(doc, titulo, level=1)
            add_paragraph(doc, secoes[chave])

    add_heading(doc, "Escopo, Investimento e Condições", level=1)
    campos = dados.get("campos_deterministicos", {})
    pendentes_encontrados = []
    for chave, titulo in CAMPO_TITULOS.items():
        valor = campos.get(chave, f"{PENDENTE_MARCADOR} — não fornecido]")
        eh_pendente = PENDENTE_MARCADOR in str(valor)
        if eh_pendente:
            pendentes_encontrados.append(titulo)
        p = doc.add_paragraph()
        p.add_run(f"{titulo}: ").bold = True
        add_paragraph_run = p.add_run(str(valor))
        if eh_pendente:
            add_paragraph_run.font.color.rgb = ALERTA_COR
            add_paragraph_run.bold = True

    templates_usados = dados.get("templates_usados", [])
    if templates_usados:
        add_heading(doc, "Nota interna — templates de referência usados", level=2)
        for t in templates_usados:
            doc.add_paragraph(t, style="List Bullet")

    if pendentes_encontrados:
        doc.add_paragraph()
        add_paragraph(
            doc,
            "ATENÇÃO: esta proposta tem campo(s) determinístico(s) pendente(s) de preenchimento "
            f"a partir da tabela oficial da organização: {', '.join(pendentes_encontrados)}. "
            "Não envie ao cliente antes de resolver.",
            alerta=True,
        )

    return doc


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("dados_json", help="JSON com os dados da proposta (ver docstring)")
    parser.add_argument("saida_docx", help="Caminho do .docx de saída")
    parser.add_argument("--template", default=None, help="Template .docx opcional para herdar estilo")
    args = parser.parse_args()

    with open(args.dados_json, encoding="utf-8") as f:
        dados = json.load(f)

    if args.template and not Path(args.template).exists():
        print(f"Aviso: template {args.template} não encontrado, gerando sem template.", file=sys.stderr)
        args.template = None

    doc = build_document(dados, args.template)
    doc.save(args.saida_docx)
    print(f"Proposta salva em: {args.saida_docx}")

    # Sinaliza claramente no console se ficou algo pendente, pra não passar despercebido
    campos = dados.get("campos_deterministicos", {})
    pendentes = [k for k, v in campos.items() if PENDENTE_MARCADOR in str(v)]
    if pendentes:
        print(
            f"ATENÇÃO: campos determinísticos pendentes: {', '.join(pendentes)}. "
            "Não considere o fluxo concluído sem resolver isso.",
            file=sys.stderr,
        )


if __name__ == "__main__":
    main()
